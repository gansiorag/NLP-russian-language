"""
modules get names of professions from doc files

start 20/12/2021
end 23/12/2021
author : Gansior Alexandr
gansior@gansior.ru  +79173383804
"""
    
import glob
import docx
import textract
import re

rull_one = re.compile(r'[0-9.-]')

def list_names(name_folder:str) -> list:
    list_files = glob.glob('/home/al/Projects_My/NLP-russian-language/NLP_gansior/professions/dataset_profession/ЕТКС_д/ЕТКС/*.docx')
    # for name in list_files:
    #     print(name)
    return list_files
    

def work_with_file(name_file,  pro_set:set) -> set:
    # open connection to Word Document
    doc = docx.Document(name_file)
    # read in each paragraph in file
    # for p in doc.paragraphs:
    #     print(p.text) 
    
    print(len(doc.tables))
    #pred_text =''
    #print(doc.tables)
    for rrow in doc.tables[-1].rows:
        for cc in rrow.cells:
            tt =rull_one.sub('',cc.text.strip())
            if tt :
                print(tt)
                pro_set.add(tt)
    #     pred_text = rrow.cells[2].text.strip()
    #     print('0 - ', end = '  ')
    #     print(rrow.cells[0].text.strip())
    #     print('1 - ',rrow.cells[1].text.strip())
    #     print('2 - ',rrow.cells[2].text.strip())
    #     print('3 - ',rrow.cells[3].text.strip())
    #     print()      
        # if rrow.cells[0].text.strip():
        #     if pred_text :
        #         print(pred_text.strip())
        #         pred_text = rrow.cells[1].text
        #     else:
        #         pred_text = rrow.cells[1].text
        # else :
        #     pred_text +=(' ' + rrow.cells[1].text)
    return pro_set


if __name__ == '__main__':
    name_folder = '/home/al/Projects_My/NLP-russian-language/NLP_gansior/professions/dataset_profession/ЕТКС_д/ЕТКС/'
    list_file = list_names(name_folder)
    k = 0
    #work_with_file('/home/al/Projects_My/NLP-russian-language/NLP_gansior/professions/dataset_profession/ЕТКС_д/ЕТКС/Выпуск 40.docx_err')
    pro_set = set()
    for name_file in list_file:
        print(name_file)
        k += 1
        pro_set = work_with_file(name_file,  pro_set)
        #if k == 3 : break
    r_f = open('/home/al/Projects_My/NLP-russian-language/NLP_gansior/professions/dataset_profession/list_prof.txt', 'w')
    for ss in pro_set:
        r_f.write(ss+'\n')
    print(len(pro_set))